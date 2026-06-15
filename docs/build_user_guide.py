from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "SupportCaseManager_UserGuide_ja.md"
OUT_PATH = ROOT / "docs" / "SupportCaseManager_UserGuide_ja.docx"


def set_run_font(run, font_name: str = "Meiryo") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False, size: int | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    set_run_font(r)
    return p


def add_cover_page(doc: Document, title: str, meta: dict[str, str]) -> None:
    title_text = title.strip() if title else "利用手順書"
    p = add_paragraph(doc, title_text, style="Heading 1", bold=True, size=24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = "超詳細版"
    p2 = add_paragraph(doc, subtitle, bold=True, size=16)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.get("date"):
        p3 = add_paragraph(doc, f"作成日: {meta['date']}", size=12)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.get("audience"):
        p4 = add_paragraph(doc, f"対象読者: {meta['audience']}", size=11)
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.get("scope"):
        p5 = add_paragraph(doc, f"対象範囲: {meta['scope']}", size=11)
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p = paragraph._p
    insert_index = 0
    for i, child in enumerate(p):
        if child.tag == qn("w:pPr"):
            insert_index = i + 1
            break
    p.insert(insert_index, start)
    p.append(end)


def add_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_manual_toc(doc: Document, headings: list[tuple[int, str, str]]) -> None:
    add_paragraph(doc, "目次", style="Heading 1")
    for level, text, anchor in headings:
        if level > 2:
            continue
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Pt(12 * (level - 1))
        add_hyperlink(p, text, anchor)
    doc.add_page_break()


def add_word_toc(doc: Document) -> None:
    add_paragraph(doc, "目次", style="Heading 1")
    p = doc.add_paragraph()
    # Word TOC field: levels 1-2, with hyperlinks and page numbers
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    doc.add_page_break()


def apply_base_styles(doc: Document, font_name: str = "Meiryo") -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Paragraph"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = font_name
            if style_name == "Normal":
                style.font.size = Pt(10.5)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def collect_headings(lines: list[str]) -> list[tuple[int, str, str]]:
    headings: list[tuple[int, str, str]] = []
    skip_title = True
    for line in lines:
        if line.startswith("# "):
            text = line[2:].strip()
            if skip_title:
                skip_title = False
                continue
            if text in ("表紙", "目次"):
                continue
            headings.append((1, text, f"toc_{len(headings) + 1}"))
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            headings.append((2, text, f"toc_{len(headings) + 1}"))
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            headings.append((3, text, f"toc_{len(headings) + 1}"))
            continue
    return headings


def main() -> None:
    if not MD_PATH.exists():
        raise SystemExit(f"Markdown not found: {MD_PATH}")

    text = MD_PATH.read_text(encoding="utf-8")
    text = text.replace("{{DATE}}", datetime.now().strftime("%Y/%m/%d"))

    doc = Document()
    apply_base_styles(doc)

    lines = [line.rstrip() for line in text.splitlines()]
    meta = {
        "date": "",
        "audience": "",
        "scope": "",
    }
    for line in lines:
        if line.startswith("作成日:"):
            meta["date"] = line.split(":", 1)[1].strip()
        elif line.startswith("対象読者:"):
            meta["audience"] = line.split(":", 1)[1].strip()
        elif line.startswith("対象範囲:"):
            meta["scope"] = line.split(":", 1)[1].strip()

    title = ""
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    headings = collect_headings(lines)
    add_cover_page(doc, title, meta)
    add_word_toc(doc)

    skip_manual_title = True
    skip_block = None
    heading_index = 0
    for line in lines:
        if not line:
            continue

        if line.startswith("作成日:") or line.startswith("対象読者:") or line.startswith("対象範囲:"):
            continue

        if line.startswith("# ") and skip_manual_title:
            skip_manual_title = False
            continue

        if line.startswith("# "):
            heading = line[2:].strip()
            if heading in ("表紙", "目次"):
                skip_block = heading
                continue
            if skip_block is not None:
                skip_block = None

        if skip_block is not None:
            continue

        # Headings
        if line.startswith("### "):
            p = add_paragraph(doc, line[4:], style="Heading 3")
            if heading_index < len(headings):
                add_bookmark(p, headings[heading_index][2], heading_index + 1)
                heading_index += 1
            continue
        if line.startswith("## "):
            p = add_paragraph(doc, line[3:], style="Heading 2")
            if heading_index < len(headings):
                add_bookmark(p, headings[heading_index][2], heading_index + 1)
                heading_index += 1
            continue
        if line.startswith("# "):
            p = add_paragraph(doc, line[2:], style="Heading 1")
            if heading_index < len(headings):
                add_bookmark(p, headings[heading_index][2], heading_index + 1)
                heading_index += 1
            continue

        # Image
        match = re.match(r"!\[.*?\]\((.+?)\)", line)
        if match:
            image_path = (ROOT / match.group(1)).resolve()
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.5))
            else:
                add_paragraph(doc, f"(画像が見つかりません: {match.group(1)})")
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            item = re.sub(r"^\d+\.\s+", "", line)
            add_paragraph(doc, item, style="List Number")
            continue

        # Bullet list
        if line.startswith("- "):
            add_paragraph(doc, line[2:], style="List Bullet")
            continue

        # Horizontal rule (---)
        if line.strip("-") == "":
            continue

        # Normal paragraph
        add_paragraph(doc, line)

    # Ensure fonts on all runs
    for p in doc.paragraphs:
        for r in p.runs:
            set_run_font(r)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
